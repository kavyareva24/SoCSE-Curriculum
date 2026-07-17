from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(r"C:\Users\Dell\Desktop\SoCSE-Curriculum")
source = root / "CourseFiles/CS-Core/CF_SemV_OperatingSystems.md"
out_file = root / "B24CI0503_Operating Systems_V_Sem.docx"

text = source.read_text(encoding="utf-8")

# Strip YAML front matter
if text.startswith("---\n"):
    parts = text.split("\n---\n", 1)
    if len(parts) == 2:
        text = parts[1]

lines = text.splitlines()

# Remove markdown formatting markers and build document content
content = []
for raw in lines:
    line = raw.rstrip()
    if not line.strip():
        continue
    line = line.replace("🟢", "").replace("🔵", "").replace("✅", "")
    line = line.replace("**", "")
    line = line.strip()

    if line.startswith("# "):
        content.append(("H1", line[2:].strip()))
    elif line.startswith("## "):
        content.append(("H2", line[3:].strip()))
    elif line.startswith("### "):
        content.append(("H3", line[4:].strip()))
    elif line.startswith("|"):
        continue
    elif line.startswith("- "):
        content.append(("BULLET", line[2:].strip()))
    elif line.startswith(">"):
        continue
    elif line.startswith("Status:"):
        continue
    elif line.startswith("course_code:") or line.startswith("title:") or line.startswith("programs:") or line.startswith("semester:") or line.startswith("category:") or line.startswith("ltpc:") or line.startswith("contact_hours_per_week:") or line.startswith("cie:") or line.startswith("see:") or line.startswith("total_marks:") or line.startswith("aicte_category:") or line.startswith("level:") or line.startswith("status:"):
        continue
    else:
        content.append(("P", line))

# Create document

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('B24CI0503 – Operating Systems')
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('V Semester Course File')
run.italic = True
run.font.size = Pt(12)

# Add content
for kind, value in content:
    if kind == 'H1':
        p = doc.add_paragraph()
        p.add_run(value).bold = True
        p.runs[0].font.size = Pt(14)
    elif kind == 'H2':
        p = doc.add_paragraph()
        p.add_run(value).bold = True
        p.runs[0].font.size = Pt(12)
    elif kind == 'H3':
        p = doc.add_paragraph()
        p.add_run(value).bold = True
        p.runs[0].font.size = Pt(11)
    elif kind == 'BULLET':
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(value)
    else:
        p = doc.add_paragraph()
        p.add_run(value)

# Save
out_file = out_file.resolve()
doc.save(out_file)
print(out_file)
